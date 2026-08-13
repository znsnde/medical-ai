"""
标准论文格式排版脚本
保留前两页不变，从第三页开始按标准格式排版
"""
import docx
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = r"F:\Python\Project\240034301015-李子旭-创新训练II报告模板.docx"
DST = r"F:\Python\Project\240034301015-李子旭-创新训练II报告模板_已排版.docx"

doc = docx.Document(SRC)

# 标准论文格式常量
FONT_BODY = "宋体"
FONT_HEADING = "黑体"
SIZE_BODY = Pt(12)       # 小四
SIZE_H1 = Pt(15)         # 小三 标题
SIZE_H2 = Pt(14)         # 四号 标题
SIZE_TITLE = Pt(18)      # 小二 论文标题
LINE_SPACING = 1.5
INDENT_FIRST = Cm(0.74)  # 首行缩进2字符

# ── 判断行类型 ──
def is_h1(text):
    return text.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、"))

def is_h2(text):
    return text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")) and len(text) < 60

def is_abstract(text):
    return text.strip() == "摘要"

def is_keywords(text):
    return text.startswith("关键词") or text.startswith("關鍵詞")

def is_title(text):
    return len(text) > 0 and len(text) < 30 and "群智能" in text

# ── 找到要排版的起始段落 ──
# 保留前两页不变（封面+目录），从摘要开始排版
start_idx = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "摘要":
        start_idx = i
        break
# 如果没找到摘要，从第53段开始
if start_idx == 0:
    start_idx = 53

print(f"从段落 {start_idx} 开始排版（摘要）")

# ── 排版 ──
for i, p in enumerate(doc.paragraphs):
    if i < start_idx:
        continue  # 保留前两页

    text = p.text.strip()
    if not text:
        continue

    # 清除原有格式
    for run in p.runs:
        run.font.name = None
        run.font.size = None
        run.font.bold = None
        run.font.italic = None
        run.font.color.rgb = None

    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 设置对齐
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── 标题格式 ──
    if is_abstract(text) or is_keywords(text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = FONT_HEADING
            run.font.size = SIZE_H1
            run.font.bold = True

    elif is_h1(text):
        # 一级标题：黑体 小三
        for run in p.runs:
            run.font.name = FONT_HEADING
            run.font.size = SIZE_H1
            run.font.bold = True

    elif is_h2(text):
        # 二级标题：黑体 四号
        for run in p.runs:
            run.font.name = FONT_HEADING
            run.font.size = SIZE_H2
            run.font.bold = True

    else:
        # 正文：宋体 小四，首行缩进
        p.paragraph_format.first_line_indent = INDENT_FIRST
        for run in p.runs:
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.bold = False

    # 设置中文字体（word中需要同时设置western字体和east asia字体）
    for run in p.runs:
        run.font.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), FONT_BODY if run.font.name == FONT_BODY else FONT_HEADING)

doc.save(DST)
print(f"排版完成！已保存至: {DST}")
